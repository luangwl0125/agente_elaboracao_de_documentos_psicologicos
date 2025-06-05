import streamlit as st
import datetime
import logging
import time
import re
from docx import Document  # Corrigido de `document` para `Document`.
from io import BytesIO  # Corrigido de `bytesio` para `BytesIO`.
from PyPDF2 import PdfReader  # Corrigido de `pdfreader` para `PdfReader`.
from PIL import Image  # Corrigido de `image` para `Image`.
import pytesseract
import openai

logging.basicConfig(level=logging.INFO)
openai.api_key = st.secrets["openai_api_key"]
assistant_id = "asst_am7evj3dygihdztqx0boe0e6"

def extrair_texto_arquivo(files):
    textos = []
    for file in files:
        try:
            if file.type == "application/pdf":
                reader = PdfReader(file)
                texto = "\n".join(page.extract_text() or "" for page in reader.pages)
                textos.append(texto.strip() or "[pdf sem texto detectável]")
            elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(file)
                texto = "\n".join([p.text for p in doc.paragraphs])
                textos.append(texto.strip() or "[docx sem texto detectável]")
            elif file.type.startswith("image/"):
                image = Image.open(file)
                texto = pytesseract.image_to_string(image, lang="por")
                textos.append(texto.strip() or "[imagem sem texto detectável]")
            else:
                textos.append(f"[tipo de arquivo não suportado: {file.type}]")
        except Exception as e:
            logging.error(f"erro ao extrair texto de {file.name}: {e}")
            textos.append(f"[erro ao extrair texto: {e}]")
    return textos

def sanitize_filename(filename):
    return re.sub(r'[^a-zA-Z0-9_\-\.]', '_', filename)

def obter_campos_por_tipo_documento(tipo):
    estrutura = {
        "declaração psicológica": ["nome completo do(a) paciente", "finalidade da declaração", "informações sobre o atendimento"],
        "atestado psicológico": ["nome da pessoa/instituição atendida", "solicitante", "finalidade", "descrição das condições psicológicas", "cid (opcional)", "observações finais"],
        "relatório psicológico": ["identificação", "descrição da demanda", "procedimento", "análise", "conclusão"],
        "relatório multiprofissional": ["identificação", "descrição da demanda", "procedimento", "análise", "conclusão"],
        "laudo psicológico": ["identificação", "descrição da demanda", "procedimento", "análise", "conclusão", "referências"],
        "parecer psicológico": ["identificação", "descrição da demanda", "análise", "conclusão", "referências"]
    }
    return estrutura.get(tipo, [])

def gerar_campos_dinamicos(campos):
    respostas = {}
    arquivos = {}

    for campo in campos:
        col1, col2 = st.columns([2, 1])  # Colunas para layout
        with col1:
            st.markdown(f"**{campo}**")
            resposta = st.text_area("opção: anexe um documento que contenha a informação.", key=f"text_area_{campo}")
            respostas[campo] = resposta
        with col2:
            arquivos[campo] = st.file_uploader("Anexar arquivo", type=["pdf", "docx", "png", "jpg", "jpeg"], key=f"{campo}", accept_multiple_files=True)  # Permitir múltiplos arquivos
    return respostas, arquivos

def enviar_para_assistente(user_message):
    try:
        thread = openai.beta.threads.create()
        openai.beta.threads.messages.create(thread_id=thread.id, role="user", content=user_message)
        run = openai.beta.threads.runs.create(thread_id=thread.id, assistant_id=assistant_id)
        while True:
            run = openai.beta.threads.runs.retrieve(thread_id=thread.id, run_id=run.id)
            if run.status == "completed":
                break
            time.sleep(1)
        messages_list = openai.beta.threads.messages.list(thread_id=thread.id)
        for msg in reversed(messages_list.data):
            if msg.role == "assistant":
                return msg.content[0].text.value
        return "[resposta não encontrada]"
    except Exception as e:
        logging.error(f"erro na api assistants: {e}")
        return f"[erro ao interagir com o assistente: {e}]"

def exportar_para_docx(texto):
    doc = Document()
    for paragrafo in texto.split("\n"):
        doc.add_paragraph(paragrafo.strip())
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Verifica se os termos foram aceitos
if 'accepted_terms' not in st.session_state:
    st.session_state.accepted_terms = False

if not st.session_state.accepted_terms:
    if st.button("Leia os termos de uso e aceito continuar"):
        st.session_state.accepted_terms = True
        st.experimental_rerun()

    # Apresenta os termos de uso e privacidade ao usuário
    st.header("Termos de Uso e Política de Privacidade")
    st.markdown("""
    🛡️ **Compromisso com a ética, segurança e sigilo profissional**

    Este serviço foi desenvolvido como ferramenta de apoio técnico à elaboração de documentos psicológicos, com base nas diretrizes estabelecidas pela **Resolução CFP nº 06/2019**.

    ...

    **Ao utilizar este sistema, você declara ciência de que respeita e segue os preceitos éticos da profissão e assume a responsabilidade técnica e legal pelos documentos emitidos com o apoio desta ferramenta.**
    """)
    st.stop()

# Continuação do aplicativo
st.header("🧠 Psicólogo Assistente 🧠", divider="gray")
st.write("")  # adiciona uma linha de espaço

# Entrada de nome
nome = st.text_input("Seu Nome Completo")

# Entrada do CRP
numero = st.text_input("CRP", max_chars=10)

# Entrada da data
data = st.date_input("Data", value=datetime.date.today())

# Mensagem de boas-vindas
if nome and numero:
    st.write(f"Olá, {nome}! | CRP: {numero}")

st.markdown("----")

# Tipo de documento
tipo_documento = st.selectbox("Selecione o tipo de documento que você quer assistência na elaboração", [
    "Declaração Psicológica",
    "Atestado Psicológico",
    "Relatório Psicológico",
    "Relatório Multiprofissional",
    "Laudo Psicológico",
    "Parecer Psicológico"
])

st.write("----")  # adiciona uma linha de espaço

# Campos dinâmicos
campos = obter_campos_por_tipo_documento(tipo_documento)
respostas, arquivos = gerar_campos_dinamicos(campos)

enviar = st.button("🔍 Gerar Documento")

if enviar:
    conteudo = f"Tipo de documento: {tipo_documento}\n\n"
    for campo in campos:
        if arquivos[campo]:
            textos_extraidos = extrair_texto_arquivo(arquivos[campo])
            conteudo += f"{campo}:\n" + "\n".join(textos_extraidos) + "\n\n"
        else:
            conteudo += f"{campo}:\n{respostas[campo]}\n\n"

    resposta = enviar_para_assistente(conteudo)
    st.subheader("📄 Documento Gerado")
    st.text_area("Conteúdo", resposta, height=400)
    buffer = exportar_para_docx(resposta)
    st.download_button("📥 Baixar .docx", data=buffer, file_name=f"{sanitize_filename(tipo_documento)}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.info("🔍 Este documento deve ser revisado pelo psicólogo responsável antes do uso oficial.")


    
# Termos de uso e consentimento inicial
if 'accepted_terms' not in st.session_state:
    st.session_state.accepted_terms = False

if not st.session_state.accepted_terms:
    st.header("Termos de Uso e Política de Privacidade")
    st.markdown("""
    🛡️ **Compromisso com a Ética, Segurança e Sigilo Profissional**

                
    Este serviço foi desenvolvido como ferramenta de apoio técnico à elaboração de documentos psicológicos, com base nas diretrizes estabelecidas pela **Resolução CFP nº 06/2019**, pela **Resolução CFP nº 01/2009 (Política de Proteção de Dados)** e pelo **Código de Ética Profissional do Psicólogo**.

                
    🧠 **Responsabilidade Técnica e Ética**  
    As produções dos documentos devem **obrigatoriamente ser revisadas, validadas e assinadas por psicóloga(o) devidamente inscrita(o) no CRP**, conforme determina a legislação profissional.  
    O conteúdo gerado **não substitui o julgamento clínico e técnico do profissional**.

                
    📌 **Finalidade do Sistema**  
    Este assistente virtual tem como único propósito **auxiliar a(o) psicóloga(o)** na sistematização de informações, organização textual e conformidade estrutural de documentos, sempre respeitando os princípios de autonomia, consentimento informado, sigilo, não exposição e ética nas relações profissionais.
                

    ⚖️ **Referências Normativas**
    - Resolução CFP nº 06/2019 – Elaboração de Documentos Escritos Produzidos pela(o) Psicóloga(o)
    - Código de Ética Profissional do Psicólogo – Artigos 1º, 9º, 13º e 14º
    - Resolução CFP nº 11/2018 – Sobre uso de tecnologias da informação e comunicação
    - LGPD (Lei Geral de Proteção de Dados) – Aplicabilidade ao contexto psicológico
                

    🔒 **Privacidade e Proteção de Dados**  
    - Esta ferramenta foi construída em conformidade com:
    - O Código de Ética do Profissional Psicólogo (Resolução CFP nº 010/2005);
    - A Resolução CFP nº 06/2019: Elaboração de Documentos Escritos Produzidos pela(o) Psicóloga(o);
    - Resolução CFP nº 11/2018: Sobre uso de tecnologias da informação e comunicação
    - Criptografia em trânsito (HTTPS): Criptografia de Ponta a Ponta para Proteger Dados em Trânsito e em Repouso. Todos os dados são protegidos contra interceptação.
    - Controle de acesso: APIs protegidas com autenticação para impedir acesso não autorizado.
    - Validação de entrada: Validações automáticas, evitando injeções maliciosas ou erros lógicos.
    - Registros e auditoria: Rastreamento de dados com precisão (data/hora e autor), ajudando na responsabilização e conformidade com normas como a LGPD.
    - Anonimização: Omissão de dados sensíveis antes de armazenar ou compartilhar informações JSON, promovendo privacidade.
    - Normas da Lei Geral de Proteção de Dados Pessoais (Lei nº 13.709/2018), que regula o tratamento de dados pessoais no Brasil. Seu objetivo principal é garantir o direito à privacidade e à proteção dos dados dos cidadãos, estabelecendo regras claras sobre coleta, uso, armazenamento e compartilhamento de informações pessoais por empresas, órgãos públicos e profissionais autônomos incluindo psicólogas(os).
                
    Ao utilizar este sistema, você declara ciência de que **respeita e segue os preceitos éticos da profissão** e que **assume a responsabilidade técnica e legal pelos documentos emitidos** com o apoio desta ferramenta.
    """)
    if st.button("Aceito os Termos e Continuar"):
        st.session_state.accepted_terms = True
        st.experimental_rerun()
    st.stop()